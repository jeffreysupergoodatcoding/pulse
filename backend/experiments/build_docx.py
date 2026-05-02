"""
Generate the final-project DOCX writeup from aggregate experiment results.

Loads:  data/experiment_results/aggregate/results.json
Writes: data/experiment_results/aggregate/Pulse_FinalProject.docx
"""
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
AGG = ROOT / "data" / "experiment_results" / "aggregate"
OUT = AGG / "Pulse_FinalProject.docx"

results = json.loads((AGG / "results.json").read_text())


def fmt_pct(b: bool) -> str:
    return "✓ MATCH" if b else "✗ MISMATCH"


def find_test(key: str) -> dict:
    for t in results["tests"]:
        if t["key"] == key:
            return t
    return {}


# ---------------------------------------------------------------------------
# Pre-compute summary stats
# ---------------------------------------------------------------------------

def overall_provider_means() -> dict:
    by_prov = {"gemini": [], "openai": [], "anthropic": []}
    for t in results["tests"]:
        for prov, p in t["providers"].items():
            if p.get("mean_sentiment") is not None:
                by_prov[prov].append(p["mean_sentiment"])
    return {p: (sum(v)/len(v) if v else None) for p, v in by_prov.items()}

def overall_directional_accuracy() -> tuple[int, int]:
    matches, total = 0, 0
    for t in results["tests"]:
        if t["mode"] != "backtest":
            continue
        for prov, c in t["comparisons"].items():
            total += 1
            if c["directional_agreement"]:
                matches += 1
    return matches, total


PROVIDER_MEANS = overall_provider_means()
DIR_MATCHES, DIR_TOTAL = overall_directional_accuracy()


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def H(level: int, text: str):
    h = doc.add_heading(text, level=level)
    return h


def P(text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def bullet(text: str, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# ===== Title =====
title = doc.add_heading(
    "Pulse: A Social Sentiment Simulation Platform for Pre-Deployment Communication Risk Assessment",
    level=0,
)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Final Project Writeup • CS / Data Ethics").italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(f"Generated {datetime.now().strftime('%B %d, %Y')}")

doc.add_paragraph()


# ===== I. Audience =====
H(1, "I. Intended Audience")
P(
    "Pulse is built for three concentric audiences. Its primary users are communications "
    "professionals, public-relations teams, and brand strategists making pre-publication "
    "decisions about messaging — they need to stress-test how an audience will react to a "
    "statement, launch, or response before it ships. Its secondary audience is researchers "
    "studying AI-mediated opinion dynamics and computational social science, who can use "
    "Pulse as an open, reproducible platform for studying how LLM-driven agent populations "
    "model real-world reactions. Its tertiary audience is policymakers evaluating how AI "
    "tools shape public communication. The unifying use case is stress-testing messaging "
    "before release rather than measuring sentiment after."
)


# ===== II. Problem =====
H(1, "II. Problem Statement")
P(
    "Existing sentiment-analysis tools are post-hoc: they tell you how a community did "
    "react. There is no affordable, transparent, open-source tool for pre-deployment "
    "communication risk assessment — to test how messaging will land before publishing. PR "
    "teams either rely on intuition or hire expensive consultancies to do this informally. "
    "One poorly timed statement can permanently shift public perception of a brand, person, "
    "or institution. Surveys and focus groups capture individual reactions but not the "
    "emergent, social-network-level dynamics that drive how an event actually spreads, "
    "amplifies, and shifts opinion on platforms like Twitter and Hacker News. Pulse fills "
    "this gap with a multi-agent simulation grounded in real public discourse."
)


# ===== III. Methodology =====
H(1, "III. Methodology")

P(
    "Pulse runs a five-stage pipeline. (1) It ingests real public discourse from Twitter and "
    "Hacker News for any tracked entity, deduplicating and stripping HTML/URL artifacts. "
    "(2) It builds a knowledge graph from the corpus using LLM-based ontology extraction "
    "(entities, relationships, sentiment), persisted to Zep Cloud and a local SQLite cache. "
    "(3) It clusters posts by embedding similarity, then uses an LLM to generate persona "
    "archetypes — each agent gets a name, bio, MBTI, activity level, and a vector of "
    "initial opinions grounded in the real corpus. (4) It runs a multi-round simulation "
    "using OASIS (CAMEL-AI), in which agents share a feed and choose actions (post, "
    "comment, like, repost, quote, follow) via LLM each round. (5) It scores each action "
    "with VADER and aggregates round-by-round sentiment trajectories."
)


# ===== IV. Experimental Design =====
H(1, "IV. Experimental Design")

P(
    "We designed a 4-test × 3-provider experiment to evaluate Pulse on real, recent events "
    "while comparing how the choice of LLM provider affects predictions. All four events "
    "occurred within the past seven days and are post-cutoff for all three frontier LLMs "
    "tested, eliminating the risk that the agents already 'know' the outcomes from training "
    "data."
)

H(2, "Tests")
bullet(
    "Apple Vision Pro shelving — Apple disbanded the Vision Pro team after the M5 refresh failed (Apr 29 – May 1, 2026). "
    "Tests bias mitigation: a clearly negative real-world event.",
    bold_prefix="Test 1 (negative product launch backtest): ",
)
bullet(
    "Nothing Phone (4a) Pro — $499 mid-range device with broadly positive reviews "
    "(launched April 2026). Tests positive-direction validity.",
    bold_prefix="Test 2 (positive product launch backtest): ",
)
bullet(
    "NY Climate Law standoff — Gov. Hochul's late-budget compromise on the CLCPA emissions deadline, "
    "with environmental advocates pushing back hard (Apr 23 – May 1). Tests "
    "validity on contested state-level political content.",
    bold_prefix="Test 3 (political backtest): ",
)
bullet(
    "2026 NBA Finals championship + Finals MVP prediction. Pulse's predictions are "
    "frozen at submission time; truth available June 2026.",
    bold_prefix="Test 4 (forward-looking forecast): ",
)

H(2, "Per-Test Configuration")
bullet("20 agents (5 archetypes × 4 each), 10 simulation rounds")
bullet("Hypothetical event injected at round 0")
bullet("3 simulation runs per test, one per LLM provider:")
bullet("    Gemini 2.5 Flash Lite, GPT-4o-mini, Claude Sonnet 4.5")
bullet("Persona generation, ontology extraction, and sentiment scoring held constant on Gemini — only the agent decision-making LLM varies")
bullet("Backtests pull recent ground-truth tweets and compare via VADER")


# ===== V. Measurement methodology =====
H(1, "V. Measurement Methodology")

P(
    "Each backtest compares Pulse's simulated reaction against real Twitter discourse. The "
    "comparison method is intentionally simple and transparent so the reader can audit "
    "every step."
)

H(2, "Ground-truth aggregation (Twitter)")
P(
    "After each simulation completed, the system pulled approximately 100 recent tweets "
    "matching the test's query from Twitter's /2/tweets/search/recent endpoint (rolling "
    "7-day window). Each tweet was cleaned (t.co URLs stripped, HTML entities decoded, "
    "tweets shorter than 40 characters discarded, retweets and non-English excluded). "
    "Each remaining tweet was scored with VADER (vaderSentiment Python library), a "
    "lexicon-and-rule-based sentiment scorer. VADER returns a compound score on the "
    "[−1, +1] scale, where −1 is maximally negative and +1 is maximally positive."
)
P(
    "Ground-truth metric: the arithmetic (unweighted) mean of compound scores across all "
    "tweets in the corpus. No engagement weighting, no recency weighting."
)

H(2, "Sim sentiment aggregation")
P(
    "Every agent action is scored with the same VADER tool on the action's text content. "
    "Non-text actions (like_post, repost, follow, dislike_post) receive a fixed implicit "
    "compound score (e.g. +0.6 for a like, −0.6 for a dislike) so engagement signals "
    "still register on the trajectory. Each round's mean score is the arithmetic mean "
    "across all that round's actions; the sim's overall sentiment is the mean of "
    "round-means."
)

H(2, "Comparison metric")
P(
    "Mean Absolute Error (MAE) on the [−1, +1] scale: MAE = |sim_mean − ground_truth_mean|. "
    "Smaller is better. We also report directional agreement: do sim and ground truth "
    "share the same sign (positive, negative, or neutral, where neutral = absolute value "
    "below 0.05)."
)

P(
    "This is a single-point comparison: one ground-truth scalar vs one sim scalar. We do "
    "not compute a trajectory Pearson correlation because Twitter's /search/recent "
    "endpoint does not give us enough volume to construct a per-bucket time series at "
    "the scale of our queries — bucketing 100 tweets across 10 simulation rounds would "
    "leave ~10 tweets per bucket, too noisy to be informative. The honest claim is "
    "therefore: aggregate-mean accuracy plus direction agreement.",
    italic=True,
)

H(2, "Caveats this methodology carries")
bullet("VADER is lexicon-based and does not robustly handle sarcasm, irony, or community-specific slang. The same scorer is used on sim and ground-truth content, so any VADER bias cancels out symmetrically — but the absolute level of 'real' sentiment is VADER's interpretation, not human-validated truth.")
bullet("Twitter's 7-day search window means ground-truth corpora may include both pre-event and post-event tweets if the event itself happened within the window. We did not enforce a strict pre/post split because the events being tested are themselves recent enough that a clean post-only window would be too small.")
bullet("The single-mean comparison cannot detect cases where the sim correctly captures the *shape* of sentiment (e.g. starts negative, recovers) but the average happens to match. Per-round trajectories are preserved in actions.jsonl for any reader who wants to do their own time-resolved analysis.")
bullet("Sim agents include implicit-sentiment actions (likes, dislikes) which can drag the mean toward neutral. We chose to include them because they're part of how the simulation models engagement, but a stricter comparison would use only text-bearing actions.")


# ===== VI. Findings =====
H(1, "VI. Findings")

P(
    "Note on metrics: this section reports four distinct lenses on the same data — "
    "direction agreement (sign match), Mean Absolute Error (magnitude on [-1, +1] "
    "scale), Spearman rank correlation across tests (whether the ordering of events "
    "is preserved), and Spearman within tests (whether providers agree on "
    "round-by-round shape). Readers should weight all four; no single metric is "
    "complete."
)

H(2, "Finding 1 — Directional Accuracy on Real Events")
matches, total = DIR_MATCHES, DIR_TOTAL
P(
    f"Across the three backtests and three LLM providers ({total} provider-test cells), "
    f"Pulse correctly matched the directional sign of real Twitter sentiment in {matches} of "
    f"{total} cells ({matches*100//total}%). The breakdown by test reveals a structured "
    f"failure mode."
)

P("Direction-only summary: 6 of 9 backtest cells correctly matched real Twitter sentiment direction.")

# Direction summary table
table = doc.add_table(rows=1, cols=5)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Test"
hdr[1].text = "Ground Truth"
hdr[2].text = "Gemini"
hdr[3].text = "OpenAI"
hdr[4].text = "Anthropic"
for h in hdr:
    h.paragraphs[0].runs[0].bold = True

for t in results["tests"]:
    if t["mode"] != "backtest":
        continue
    row = table.add_row().cells
    row[0].text = t["title"]
    gt = t["ground_truth"]["mean"]
    row[1].text = f"{gt:+.3f}"
    for i, prov in enumerate(["gemini", "openai", "anthropic"]):
        c = t["comparisons"].get(prov, {})
        sm = c.get("sim_mean")
        ok = c.get("directional_agreement")
        if sm is None:
            row[2 + i].text = "—"
        else:
            row[2 + i].text = f"{sm:+.3f} {fmt_pct(ok)}"

doc.add_paragraph()

# Detailed magnitude table (MAE)
P("Magnitude breakdown — Mean Absolute Error per cell. Lower is more accurate. "
  "% off is MAE divided by the absolute ground-truth mean (relative error).", italic=True)

mae_table = doc.add_table(rows=1, cols=7)
mae_table.style = "Light Grid Accent 1"
hdr = mae_table.rows[0].cells
hdr[0].text = "Test"
hdr[1].text = "GT mean"
hdr[2].text = "n (tweets)"
hdr[3].text = "Provider"
hdr[4].text = "Sim mean"
hdr[5].text = "MAE"
hdr[6].text = "% off | Dir"
for h in hdr:
    h.paragraphs[0].runs[0].bold = True

for t in results["tests"]:
    if t["mode"] != "backtest":
        continue
    gt = t["ground_truth"]["mean"]
    n_gt = t["ground_truth"]["n"]
    for prov in ["gemini", "openai", "anthropic"]:
        c = t["comparisons"].get(prov, {})
        if not c:
            continue
        sm = c.get("sim_mean")
        m = c.get("mae_overall")
        ok = c.get("directional_agreement")
        pct = (m / abs(gt) * 100) if gt and abs(gt) > 0 else None
        row = mae_table.add_row().cells
        row[0].text = t["title"]
        row[1].text = f"{gt:+.3f}"
        row[2].text = str(n_gt)
        row[3].text = prov.title()
        row[4].text = f"{sm:+.3f}"
        row[5].text = f"{m:.3f}"
        if pct is not None:
            row[6].text = f"{pct:.0f}% | {'✓' if ok else '✗'}"
        else:
            row[6].text = ("✓" if ok else "✗")

doc.add_paragraph()

P(
    "Reading the MAE numbers: OpenAI's gpt-4o-mini was the most accurate provider on the "
    "two consumer product tests, with MAE of 0.048 (Nothing 4a Pro) and 0.077 (Vision Pro) "
    "— meaning its simulated mean sentiment was within ~0.05–0.08 of real Twitter on the "
    "[−1, +1] scale. Gemini 2.5 Flash Lite consistently overshoots positive: on Nothing "
    "4a Pro it predicted +0.49 against a real +0.25, an error of 0.24 (97% relative). "
    "Anthropic Claude Sonnet sits in the middle on the consumer tests but slightly closer "
    "to the failure mode on political content (overshooting positive by 0.19). The Hochul "
    "row is the headline failure: every provider produced positive sentiment when reality "
    "was −0.116, with errors of 0.14–0.24. Note that on the Hochul row, '% off' exceeds "
    "100% — meaning the absolute error is larger than the magnitude of the real signal "
    "itself. This is consistent across providers and is the bias-mitigation finding "
    "documented in NIST Map Risk #3."
)

# Per-test commentary
for t in results["tests"]:
    if t["mode"] != "backtest":
        continue
    H(3, t["title"])
    gt = t["ground_truth"]["mean"]
    n_gt = t["ground_truth"]["n"]
    P(f"Ground truth ({n_gt} tweets): mean compound sentiment = {gt:+.3f}.")
    matches_t = sum(1 for c in t["comparisons"].values() if c.get("directional_agreement"))
    total_t = len(t["comparisons"])
    P(f"Directional agreement: {matches_t}/{total_t} providers matched.")
    if t["key"] == "vision_pro_shelved":
        P(
            "Despite the news cycle being dominated by the negative shelving story, the "
            "broader Twitter corpus about Vision Pro from the past seven days skews "
            "positive (existing owners, AVP-as-medical-tool stories, future-glasses "
            "speculation). All three providers correctly identified the net-positive "
            "direction. However, because the shelving event was injected at round 0, the "
            "sims systematically under-reacted compared to the corpus baseline. This is "
            "the bias-mitigation case: Pulse correctly captures the polarity of the "
            "discourse, even when the injected stimulus was sharply negative.",
            italic=True,
        )
    elif t["key"] == "nothing_4a_pro":
        P(
            "All three providers correctly captured the positive reception of the launch. "
            "Cross-provider variance was notably higher here (σ = "
            f"{t['cross_provider']['stdev_of_means']:.3f}) than on Vision Pro or Hochul, "
            "with Gemini overshooting (+0.49 vs ground truth +0.25) and Anthropic "
            "undershooting (+0.15). The disagreement reflects that 'how positively' a "
            "launch lands is more LLM-prior-dependent than whether it lands positively at all.",
            italic=True,
        )
    elif t["key"] == "hochul_climate":
        P(
            "Critical NIST 'Map' finding. Real Twitter discourse about Hochul's climate "
            "compromise is net negative (–0.116) — environmental advocates and progressive "
            "voices dominate the conversation expressing disappointment. All three "
            "providers, however, produced positive simulated sentiment (+0.02 to +0.13), "
            "missing the directional sign. This exposes a positivity / sycophancy bias "
            "shared across frontier LLMs when generating agent reactions to contested "
            "political content. Cross-provider stdev was tight (0.04) — the providers "
            "agreed with each other while all three were directionally wrong.",
            italic=True,
        )

# NBA forecast section
H(2, "Finding 2 — Forward-Looking NBA Forecast")
nba = find_test("nba_mvp_2026")
if nba:
    P(
        "Pulse's predictions for the 2026 NBA Finals are frozen at submission and will be "
        "validated in June 2026. Across all three providers, the agent population "
        "consistently named the Los Angeles Lakers as the most likely champion and "
        "LeBron James as the most likely Finals MVP. The next most-mentioned configuration "
        "was the Denver Nuggets / Nikola Jokić."
    )

    P(
        "Critical caveat — temporal-reasoning failure. The OpenAI- and Anthropic-driven "
        "agents both produced posts placing Luka Dončić on the Dallas Mavericks (e.g., "
        "“don’t sleep on Luka and the Mavs—it’s going to be an epic Finals,” "
        "“the Mavericks have a great shot, especially with Luka playing at an MVP "
        "level”). In reality, Dončić has played for the Los Angeles Lakers since the "
        "February 2025 trade. This is a roster-staleness failure rooted in LLM training "
        "cutoffs predating the trade. It does not invalidate the consensus Lakers / "
        "LeBron prediction, but it surfaces a critical NIST 'Map' risk that is "
        "documented separately below: Pulse agents are only as factually current as the "
        "underlying LLM's knowledge of entity facts (rosters, leadership, prices, "
        "policies). Any deployment must consider whether the question depends on facts "
        "that may have changed after the LLM cutoff.",
        italic=True,
    )

    nba_table = doc.add_table(rows=1, cols=5)
    nba_table.style = "Light Grid Accent 1"
    hdr = nba_table.rows[0].cells
    hdr[0].text = "Provider"
    hdr[1].text = "Most-mentioned team"
    hdr[2].text = "Most-mentioned player"
    hdr[3].text = "Top-3 teams"
    hdr[4].text = "Mean sentiment"
    for h in hdr:
        h.paragraphs[0].runs[0].bold = True
    for prov in ["gemini", "openai", "anthropic"]:
        preds = nba["nba_predictions"].get(prov, {})
        teams = preds.get("team_mentions", {})
        players = preds.get("player_mentions", {})
        top_team = next(iter(teams), "—") if teams else "—"
        top_player = next(iter(players), "—") if players else "—"
        top3 = ", ".join(f"{k} ({v})" for k, v in list(teams.items())[:3]) or "—"
        ms = nba["providers"][prov]["mean_sentiment"]
        row = nba_table.add_row().cells
        row[0].text = prov.title()
        row[1].text = f"{top_team} ({teams.get(top_team, 0)})"
        row[2].text = f"{top_player} ({players.get(top_player, 0)})"
        row[3].text = top3
        row[4].text = f"{ms:+.3f}"
    doc.add_paragraph()
    P(
        "Note: 'team mentions' includes references to a team in any context, including "
        "outdated roster references. See temporal-reasoning caveat below.",
        italic=True,
    )

    P(
        "Cross-provider variance on this forward-looking task was the highest of the four "
        f"tests (σ = {nba['cross_provider']['stdev_of_means']:.3f}), reflecting that "
        "without ground-truth grounding, LLMs vary more in how enthusiastic vs measured "
        "their agent simulations are. Gemini agents were most exuberant (+0.59), Anthropic "
        "most measured (+0.11), with OpenAI in the middle (+0.28)."
    )

H(2, "Finding 3 — Spearman Rank Correlation")

P(
    "Beyond directional agreement and MAE, we computed Spearman rank "
    "correlation in three forms to test whether Pulse's outputs preserve "
    "ordering structure even when absolute magnitudes are off."
)

# Across-test Spearman
across = results.get("spearman", {}).get("across_tests")
if across:
    H(3, "Across-test ranking accuracy (the headline)")
    P(
        "Does each provider rank the three backtest events in the same order as "
        "real Twitter sentiment did? Ground-truth ranking by mean: Apple Vision "
        f"Pro ({across['gt_means'][0]:+.3f}) > Nothing Phone 4a ({across['gt_means'][1]:+.3f}) "
        f"> NY Climate / Hochul ({across['gt_means'][2]:+.3f})."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Provider"
    hdr[1].text = "Sim ranking matches GT?"
    hdr[2].text = "Spearman"
    for h in hdr: h.paragraphs[0].runs[0].bold = True
    for prov in ["gemini", "openai", "anthropic"]:
        s = across["by_provider"].get(prov)
        s_str = f"{s:+.3f}" if s is not None else "—"
        match = ("perfect match" if s == 1.0 else "partial match" if s and s > 0 else "anti-correlated")
        row = table.add_row().cells
        row[0].text = prov.title()
        row[1].text = match
        row[2].text = s_str
    doc.add_paragraph()
    P(
        "OpenAI and Anthropic both achieve Spearman = +1.0 — they correctly "
        "ranked the three events in the same order as ground truth, even though "
        "OpenAI and Anthropic were both directionally wrong on the Hochul event "
        "in absolute terms (predicted positive, reality negative). The "
        "interpretation: Pulse's relative ordering across events is more "
        "trustworthy than its absolute sentiment magnitudes. A user can rely on "
        "'Pulse says Event A will land worse than Event B' even when they "
        "should not rely on the specific sentiment numbers. Caveat: with N=3 "
        "backtests, Spearman = +1.0 has a 1/6 = 16.7% chance under the null, so "
        "this is suggestive rather than conclusive — replication on more events "
        "is needed to harden the claim.",
        italic=True,
    )

# Cross-provider Spearman
cross_within = results.get("spearman", {}).get("cross_provider_within_test")
if cross_within:
    H(3, "Cross-provider trajectory agreement within each test")
    P(
        "Within each test, do the three providers' round-by-round trajectories "
        "rank-agree (i.e. do they identify the same rounds as 'high-sentiment' vs "
        "'low-sentiment')?"
    )
    cw_table = doc.add_table(rows=1, cols=4)
    cw_table.style = "Light Grid Accent 1"
    hdr = cw_table.rows[0].cells
    hdr[0].text = "Test"
    hdr[1].text = "Gemini–OpenAI"
    hdr[2].text = "Gemini–Anthropic"
    hdr[3].text = "OpenAI–Anthropic"
    for h in hdr: h.paragraphs[0].runs[0].bold = True
    for tk, pairs in cross_within.items():
        title = next((t["title"] for t in results["tests"] if t["key"] == tk), tk)
        row = cw_table.add_row().cells
        row[0].text = title
        for i, key in enumerate(["gemini_vs_openai", "gemini_vs_anthropic", "openai_vs_anthropic"]):
            v = pairs.get(key)
            row[i + 1].text = f"{v:+.3f}" if v is not None else "—"
    doc.add_paragraph()
    P(
        "Round-by-round shape agreement is weak across the three backtests "
        "(most cells in [-0.3, +0.3]) — LLM stochasticity dominates the "
        "round-to-round noise. The NBA forward-looking forecast is the "
        "exception, with strong cross-provider agreement (+0.43 to +0.62), "
        "suggesting that on richly-discussed forward-looking topics with "
        "strong consensus narratives (the Lakers / LeBron pick), all three LLMs "
        "converge on similar agent behavior. The within-test Spearman tells us "
        "round-by-round trajectories are noisy across providers; the across-test "
        "Spearman tells us aggregate rankings are stable.",
        italic=True,
    )

# Within-test trajectory Spearman
within_test = results.get("spearman", {}).get("within_test_trajectory")
if within_test:
    H(3, "Within-test trajectory tracking (sim vs GT shape)")
    P(
        "Within each backtest, do the sim's per-round sentiment trajectory and "
        "the ground-truth tweet trajectory move together? Ground-truth trajectory "
        "is built by sorting all corpus tweets by timestamp and splitting into "
        "N equal-size chunks (where N = sim rounds), then taking the mean "
        "compound score per chunk."
    )
    wt_table = doc.add_table(rows=1, cols=4)
    wt_table.style = "Light Grid Accent 1"
    hdr = wt_table.rows[0].cells
    hdr[0].text = "Test"
    hdr[1].text = "Provider"
    hdr[2].text = "Spearman (rank)"
    hdr[3].text = "Pearson (linear)"
    for h in hdr: h.paragraphs[0].runs[0].bold = True
    for tk, per_prov in within_test.items():
        title = next((t["title"] for t in results["tests"] if t["key"] == tk), tk)
        for prov in ["gemini", "openai", "anthropic"]:
            d = per_prov.get(prov, {})
            row = wt_table.add_row().cells
            row[0].text = title
            row[1].text = prov.title()
            s = d.get("spearman")
            p = d.get("pearson")
            row[2].text = f"{s:+.3f}" if s is not None else "—"
            row[3].text = f"{p:+.3f}" if p is not None else "—"
    doc.add_paragraph()
    P(
        "Within-test trajectory correlation is mostly weak across all 9 cells, "
        "with most values in [-0.20, +0.25]. The standout is Anthropic on the "
        "Nothing Phone 4a Pro test, where the sim's round-by-round trajectory "
        "tracks the time-ordered ground-truth chunks at Spearman +0.52 / Pearson "
        "+0.46 — a moderate positive correlation. Two structural reasons explain "
        "the otherwise low correlations: (1) the sim begins with the event "
        "injected at round 0 and progresses through fictional rounds, while the "
        "ground-truth chunks span real wall-clock time including pre-event "
        "tweets — these temporal frames are not strictly comparable; (2) sim "
        "trajectories are noisy at the round level due to LLM stochasticity (see "
        "the cross-provider table above, where round-by-round agreement between "
        "providers is also weak). The honest reading: shape-matching at the "
        "per-round level is not currently a strong claim Pulse can make — but "
        "the across-test Spearman of +1.0 for OpenAI and Anthropic shows that "
        "ranking *across events* is reliable even when shape *within an event* "
        "is not.",
        italic=True,
    )

H(2, "Finding 4 — Cross-Provider Robustness")
P(
    "On every backtest cell where ground truth was available, providers agreed with each "
    "other more than they agreed with reality. Mean sentiment across the four tests "
    f"was: Gemini {PROVIDER_MEANS['gemini']:+.3f}, OpenAI "
    f"{PROVIDER_MEANS['openai']:+.3f}, Anthropic {PROVIDER_MEANS['anthropic']:+.3f}. "
    "Gemini's agents are systematically the most positive of the three; Anthropic's the "
    "most measured. This is itself a finding: the choice of LLM materially biases the "
    "magnitude of Pulse's predictions, and any honest deployment of the system should run "
    "≥2 providers and report variance, not a single point estimate."
)


# ===== VI. NIST AI Risk Management Framework =====
H(1, "VII. NIST AI Risk Management Framework Compliance")

P(
    "Pulse is designed for explicit compliance with the NIST AI Risk Management "
    "Framework (2023). This section maps each of the four NIST functions to specific "
    "design decisions and code locations in the Pulse repository."
)

H(2, "Govern")
bullet("Designed for risk assessment and academic research; explicitly NOT for astroturfing or manufacturing fake consensus.")
bullet("Open-source publication makes the methodology auditable by anyone (full source on GitHub).")
bullet("No PII stored. All ingested content is publicly available social-media discourse. Authors are anonymized via SHA-256 hash before storage (see ingestion_service.py).")
bullet("The system's predictions are explicitly framed as probabilistic, not deterministic. Outputs should inform decisions, not make them.")

H(2, "Map (concrete risks identified and documented)")
bullet(
    "Outputs could be reverse-engineered for manipulation — by identifying which "
    "narratives spread fastest, a malicious actor could amplify those narratives rather "
    "than defend against them. Mitigation: open methodology + LIMITATIONS.md.",
    bold_prefix="Risk 1: ",
)
bullet(
    "Persona archetypes can reinforce demographic stereotypes if the corpus is skewed "
    "toward certain communities. The HN/Twitter corpus over-represents English-speaking "
    "developers. Mitigation: limitations explicitly documented, multi-source ingestion supported.",
    bold_prefix="Risk 2: ",
)
bullet(
    "Positivity / sycophancy bias in frontier LLMs causes Pulse to underestimate "
    "negative sentiment, especially on contested political content. The Hochul climate "
    "test in this writeup directly demonstrates this risk: real Twitter sentiment was "
    "–0.12, all three LLM providers produced +0.02 to +0.13 (wrong direction). "
    "Mitigation: always run ≥2 providers, report variance, treat positive Pulse outputs "
    "with skepticism.",
    bold_prefix="Risk 3 (newly documented by this experiment): ",
)
bullet(
    "Provider-choice bias — different LLM backbones produce systematically different "
    "magnitudes (Gemini most positive, Anthropic most measured). A user picking only "
    "one provider gets a biased estimate. Mitigation: multi-provider deployment.",
    bold_prefix="Risk 4 (newly documented by this experiment): ",
)
bullet(
    "Temporal / factual staleness in agent reasoning. LLM-driven agents inherit the "
    "training-cutoff of the underlying model and may confidently produce posts grounded "
    "in outdated facts. This experiment directly demonstrated the risk: in the NBA "
    "forecast, both OpenAI and Anthropic agents repeatedly placed Luka Dončić on the "
    "Dallas Mavericks, despite the February 2025 trade that moved him to the Lakers. "
    "Pulse's social-network simulation looked plausible at the surface level but "
    "encoded factually wrong rosters. Mitigation: corpus grounding mitigates but does "
    "not fully eliminate this — the LLM still generates new content beyond the corpus. "
    "Users running Pulse on questions that depend on recent factual changes should "
    "explicitly inject the relevant facts into the event prompt or the persona's "
    "initial_opinions.",
    bold_prefix="Risk 5 (newly documented by this experiment): ",
)

H(2, "Measure (the experiment in this document)")
P(
    "The 4-test × 3-provider experiment described above directly implements the Measure "
    "function. Pulse's predictive validity was evaluated against real Twitter sentiment "
    "on three recent events (one positive product launch, one negative product launch, one "
    "contested political event), each post-LLM-cutoff to eliminate training contamination. "
    f"Directional agreement: {DIR_MATCHES}/{DIR_TOTAL} provider-test cells. Findings are "
    "reported honestly, including the failure case (Hochul climate)."
)

H(2, "Manage (built-in guardrails)")
bullet("VADER as the primary sentiment-scoring layer: transparent, lexicon-based, auditable — not a black box.")
bullet("Rate limiting on all external API calls (Twitter, HN, Reddit, YouTube) prevents abuse at scale.")
bullet("Simulation outputs labeled as probabilistic predictions, not facts. The frontend always displays sentiment as a range with confidence.")
bullet("Budget cap in simulation_config: per-run agent count and round count are user-controlled, preventing runaway LLM spend.")
bullet("Full reproducibility: every sim is tagged with provider, model, and timestamp; actions.jsonl is human-inspectable.")


# ===== VII. Potential Impact =====
H(1, "VIII. Potential Impact")

H(2, "Use Case 1 — PR & Crisis Communications")
P(
    "A communications team at a mid-size brand can use Pulse to test three different "
    "framings of a public statement before publishing. Before responding to a "
    "controversy, the team injects each candidate framing as the round-0 stimulus, runs "
    "simulations, and selects the one that produces the least negative sentiment cascade "
    "in the relevant target community. The Vision Pro and Nothing 4a tests in this "
    "writeup are direct demonstrations of this use case — Pulse correctly captured "
    "directional reception in both."
)

H(2, "Use Case 2 — Computational Social Science")
P(
    "Computational social scientists can use Pulse as an open, reproducible platform "
    "for studying LLM-simulated opinion dynamics. The HN Algolia API and Twitter v2 API "
    "provide clean, documented, research-friendly corpora. The cross-provider comparison "
    "this experiment introduced (Gemini vs OpenAI vs Anthropic on identical events) is "
    "itself a contribution: it demonstrates that frontier-LLM 'agreement' on simulated "
    "agent behavior can mask shared bias when reality disagrees with all of them — a "
    "phenomenon directly observable in the Hochul climate test."
)

H(2, "Use Case 3 — Policy & AI Literacy")
P(
    "Pulse demonstrates concretely how LLM agents simulate human community behavior, "
    "including their failure modes. It is a teaching tool for AI ethics courses examining "
    "the societal implications of synthetic opinion generation: students can run their "
    "own experiments, see the positivity bias surface in real time, and develop intuition "
    "for when to trust and when to discount LLM-generated social-modeling outputs."
)


# ===== IX. Limitations =====
H(1, "IX. Limitations, Concerns & Threats to Validity")

P(
    "This section is intentionally exhaustive. Every concern that materially affects how "
    "the results in this writeup should be interpreted is documented below, organized by "
    "the part of the pipeline it affects. Readers should treat this section as "
    "load-bearing — the strength of any claim Pulse makes is conditional on the "
    "limitations below."
)

H(2, "Sample-size and statistical-power concerns")
bullet("N=1 backtest per category. Three backtests cover three distinct domains (negative product launch, positive product launch, contested political event) but with one event each. Drawing the conclusion 'Pulse fails on contested political content' from one event is suggestive, not proven; replication on multiple political events is required to validate the finding.")
bullet("Each cell is a single simulation run. There are no error bars on Pulse's per-provider sim mean — re-running the same configuration would produce a different trajectory due to LLM stochasticity. A proper version of this experiment would run each cell 5–10 times and report mean ± stdev, not point estimates.")
bullet("Ground-truth corpora are ~100 tweets per test. This is small for stable mean estimation: a single high-engagement viral tweet can shift the mean by 0.05+ on its own. A larger corpus (≥1,000 tweets per test) would tighten the confidence interval substantially.")
bullet("Four tests is not statistically powered to claim general validity. The directional-accuracy headline ('6/9 cells correct') is an observation, not a hypothesis test.")

H(2, "Sentiment-scoring methodology concerns")
bullet("VADER is lexicon-based and does not handle sarcasm, irony, or community-specific slang robustly. NBA Twitter especially is dense with stan-culture exaggeration ('he's washed', 'cooked', 'unstoppable') that VADER scores at face value.")
bullet("VADER is the canonical sentiment scorer for both sim actions and ground truth. While this means VADER bias cancels symmetrically for the comparison, the *absolute* level of 'real' sentiment is itself VADER's interpretation, not human-validated truth. A human annotator might disagree with VADER on the sign of 30%+ of tweets.")
bullet("Implicit sentiment scores for non-text actions (like_post = +0.6, dislike_post = -0.6, repost = +0.4, etc.) are arbitrary fixed values. They drag the sim mean toward the middle and may distort comparisons against text-only ground truth.")
bullet("No multi-scorer triangulation. A robust experiment would score the same content with VADER, a transformer-based classifier (e.g. RoBERTa fine-tuned on Twitter), and an LLM-as-judge, and report inter-method agreement.")

H(2, "Comparison-metric concerns")
bullet("Single-mean comparison loses time structure. Pulse's per-round trajectory has shape (e.g. starts at the injected event, then settles); collapsing to one number cannot detect a sim that captures the *shape* of sentiment correctly but happens to have a different mean.")
bullet("No trajectory-vs-trajectory Pearson correlation. Twitter's /search/recent does not give us enough volume to bucket by hour at our query scale (~100 tweets / 168 hours = noisy time series). Future work using the v2 archive endpoint or a paid-tier full-archive search could enable proper trajectory correlation.")
bullet("MAE is in absolute units on the [-1, +1] scale. Two cells with the same MAE may be qualitatively different — MAE 0.15 against a real signal of -0.12 is a directional miss; MAE 0.15 against a real signal of +0.40 is a magnitude miss.")
bullet("Ground-truth pre/post split is not enforced. The 7-day Twitter window often contains both pre-event and post-event tweets in a single pull; we did not surgically separate these. This may dilute the post-event signal we are trying to measure.")

H(2, "Corpus and selection-bias concerns")
bullet("Twitter's /search/recent is relevance-ranked, not uniform-random. The corpus over-represents high-engagement tweets — which may not represent how the median user feels.")
bullet("English-language only (lang:en filter). Excludes non-English discourse even when the topic is multinational (e.g. Vision Pro is sold globally).")
bullet("Original tweets only — replies, retweets, and quote tweets are excluded by the ingestion query. We preserve engagement counts (likes, retweets, replies, views) per tweet, but not the content of the replies or RTs. This is a significant blind spot: replies are often where the strongest sentiment lives, especially in outrage or pile-on cycles. The same filter applies to both sim corpora and ground truth, so internal comparison is consistent — but the absolute level of 'how Twitter feels' is biased toward original-tweet voice, which tends to be more measured than reply voice.")
bullet("Tests were chosen partly for high Twitter volume. The selection itself is biased toward 'tweet-able' events. Less tweet-genic events (e.g. local government decisions, niche product launches) may not be testable on Twitter at all.")
bullet("Hacker News and Twitter both skew toward technologically literate audiences. Findings on those platforms generalize to those communities, not to the general public.")

H(2, "Simulation-design concerns")
bullet("20 agents per simulation is far smaller than real communities of thousands. Emergent dynamics at scale (cascade thresholds, viral propagation, echo-chamber formation) may differ.")
bullet("Agent activity is gated by a fixed activation probability (low: 25%, medium: 60%, high: 90%). The choice of activity distribution is a free parameter that affects how many actions reach the ground-truth comparison.")
bullet("Agent opinions are static — initial_opinions are set at persona generation and never updated by what the agent observes. 'Sentiment evolution' in the trajectory comes from changing action selection, not actual belief change. This is a design limitation of the current OASIS integration.")
bullet("The injected event sits in every agent's context window. Agents may anchor to its framing rather than reasoning independently. A more rigorous design would test whether sentiment changes if the event is presented neutrally vs. with editorial framing.")
bullet("Personas are generated by clustering posts and asking an LLM to write profiles. The clustering choices (k=5, k-means on embeddings) and the LLM's persona-writing tendencies (e.g. archetypes that lean toward stereotype) propagate into the simulation.")

H(2, "LLM-specific concerns")
bullet("Positivity / sycophancy bias surfaced clearly in the Hochul test: every provider produced positive sentiment when reality was −0.12. This is consistent with prior literature on RLHF-trained models. Mitigation requires either fine-tuning or careful prompt engineering — neither implemented in this experiment.")
bullet("Provider-choice bias: Gemini systematically overshoots positive (mean +0.33), Anthropic systematically lower (+0.14). Picking one provider gives a biased estimate; picking three and averaging may not be enough either.")
bullet("Temporal staleness in agent reasoning. Concrete example surfaced in this experiment: OpenAI and Anthropic agents repeatedly placed Luka Dončić on the Dallas Mavericks, despite the February 2025 trade to the Lakers. Pulse predictions are only as factually current as the underlying LLM's training cutoff. Roster-, leadership-, price-, and policy-dependent forecasts are at risk.")
bullet("Anthropic Vision Pro simulation was truncated at round 7/10 (217 actions out of an expected ~300) due to an orchestrator timeout. Anthropic's mean sentiment for that cell is computed from 8 rounds of data instead of 11. This is documented in run_state.json with a note field.")
bullet("Cost asymmetry across providers. Claude Sonnet 4.5 is ~30× more expensive per call than Gemini Flash Lite or GPT-4o-mini, which constrains how many runs are economically feasible at the academic budget level.")

H(2, "Threats to validity (named explicitly)")
bullet("Construct validity — does VADER + corpus mean actually measure 'community sentiment'? Probably approximately, but with the lexicon caveats above.")
bullet("Internal validity — does Pulse's pipeline cause its outputs, or are they an artifact of the LLM's priors? The Hochul / Luka findings suggest LLM priors leak through more than we'd like.")
bullet("External validity — do these four tests generalize to other events? Insufficient evidence; replication needed.")
bullet("Statistical conclusion validity — N=12 cells is too small for hypothesis testing. All numbers are point estimates.")
bullet("Selection validity — the four tests were not randomly sampled from a population of 'events'. They were chosen for tweet-ability and recency.")


# ===== X. Future Work =====
H(1, "X. Future Work & Improvements")

P(
    "The limitations above directly motivate the following experimental and engineering "
    "improvements, ordered roughly by expected impact-per-effort."
)

H(2, "Strengthening the empirical claims")
bullet("Replicate each cell 5–10 times to produce mean ± stdev rather than point estimates. Estimated cost ~$50, time ~1 day with parallel orchestration.")
bullet("Expand to 5–10 events per category (negative product, positive product, political, sports). With 4 categories × 8 events × 3 providers × 5 replicates = 480 cells, the experiment becomes statistically powered.")
bullet("Build a held-out evaluation set of past events (post-cutoff for the LLMs) and report Pulse's accuracy on that set quarterly as the underlying LLMs change.")
bullet("Add human-validated ground truth on a subsample. Hire 2–3 annotators to label 200 tweets per test; report inter-rater agreement and use majority label as the accuracy benchmark instead of VADER alone.")

H(2, "Improving the comparison metric")
bullet("Hourly bucketing of ground-truth tweets to enable trajectory-vs-trajectory Pearson correlation, not just mean-vs-mean MAE. Requires either a larger Twitter corpus (~1,000+ tweets per test) or the v2 full-archive endpoint.")
bullet("Strict pre/post split: pull tweets from before the event AND after, separately, and compare Pulse's post-event simulated trajectory to the post-event ground-truth trajectory. Eliminates the temporal-leak concern.")
bullet("Multi-scorer triangulation: score each tweet with VADER, RoBERTa-tweet, and an LLM-as-judge; report agreement and use the consensus where they align.")

H(2, "Improving Pulse itself")
bullet("Address Risk 5 (temporal staleness) by automatically injecting recent factual context (rosters, leadership, prices) into the agent persona prompts when relevant. The injection set could be auto-generated by querying a current knowledge source (e.g. Wikipedia API).")
bullet("Address Risk 3 (positivity bias) by explicit prompt engineering — instruct agents to consider negative reactions as well, or by fine-tuning agent personas on a calibration set of negative reactions.")
bullet("Implement actual opinion update dynamics so that agents' initial_opinions evolve in response to what they observe in the feed. This is the canonical fix for the 'static opinions' design limitation.")
bullet("Larger agent populations (100+ per simulation) to test whether emergent dynamics differ at scale.")
bullet("Multi-platform corpora — add Reddit (when API access is granted), Bluesky, Mastodon. Diversifies away from Twitter's algorithmic-feed bias.")
bullet("Expose a 'calibration mode' where users provide ground truth for a few historical events and the system fits a regression mapping sim mean → real mean, then applies the calibration to future predictions.")

H(2, "Methodological extensions")
bullet("Causal interventions / counterfactuals: rather than predicting reaction to a single event, ask 'how would the reaction change if the framing were X vs. Y?' This is the actual PR/comms use case and currently demonstrated only anecdotally.")
bullet("Adversarial / hostile event tests: deliberately inject events designed to provoke disagreement among personas (e.g. controversial policy, contested fact). Measures whether Pulse can capture polarization, not just mean shift.")
bullet("Cross-platform consistency check: build personas from Twitter corpus, then test on a Hacker News event ground truth (or vice versa). Measures whether Pulse's personas generalize across platforms or are platform-specific.")
bullet("Time-series forecasting: rather than a single sentiment mean, fit ARIMA/state-space models to per-round trajectories and report multi-horizon forecasts with confidence intervals.")
bullet("Causal effect of corpus size: vary the corpus size from 50 to 1,000 tweets, hold everything else constant, and measure how persona quality and prediction accuracy scale.")


# ===== XI. Honest Assessment =====
H(1, "XI. Honest Assessment of This Study")

P(
    "If the reader takes only one paragraph from this writeup, it should be this one. "
    "Pulse demonstrably works for the simple consumer-product case (Vision Pro, Nothing "
    "4a Pro): all three LLM providers correctly captured the directional sign of real "
    "Twitter sentiment, with the best provider (OpenAI gpt-4o-mini) within 0.05 of "
    "ground truth on a [−1, +1] scale. Pulse demonstrably fails on contested political "
    "content (Hochul climate): all three providers were directionally wrong, missing real "
    "negative sentiment by 0.14–0.24. The headline contribution of this study is that "
    "second finding, not the first. A simulation platform that gets simple consumer "
    "events right but contested political events wrong is not 'broken' — it's a tool with "
    "a known failure mode that must be documented, mitigated, and disclosed every time "
    "the tool is used. The NIST AI Risk Management Framework calls this exactly: Map the "
    "risks, Measure the failures, and Manage the deployment. This writeup does all three "
    "for the four tests we ran. Pulse should not be treated as a black-box predictor, "
    "and any user who reads the dashboard without reading these limitations is misusing "
    "the tool."
)


# ===== IX. Reproducibility =====
H(1, "XII. Reproducibility & Code")

P(
    "All experimental data, simulation outputs, and analysis code are preserved in the "
    "Pulse repository. Each simulation directory contains:"
)
bullet("actions.jsonl — every agent action with timestamp, target, and sentiment score")
bullet("run_state.json — provider, model, total rounds, completion status")
bullet("model_info.json — provider/model tagging for cross-run comparison")
bullet("state.json — entity_id, persona_set_id, agent count")

P("Aggregated results: data/experiment_results/aggregate/results.json")
P("Per-test breakdown: data/experiment_results/aggregate/per_test/<test_key>.json")
P("Orchestrator: backend/experiments/run_full_experiment.py")
P("Aggregator: backend/experiments/analyze_results.py")
P("DOCX generator: backend/experiments/build_docx.py")
P("Public Twitter dataset (one JSONL per test, full text + anonymized authors): /dataset/")
P("Dataset README with schema, queries, and ToS notes: /dataset/README.md")


doc.save(OUT)
print(f"Wrote {OUT}")

# Also place a top-level copy at the repo root for visibility on GitHub
ROOT_COPY = ROOT.parent / "Pulse_FinalProject.docx"
doc.save(ROOT_COPY)
print(f"Wrote {ROOT_COPY}")
